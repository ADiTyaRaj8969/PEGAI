#!/usr/bin/env python
"""Builds the submission document (DOCX + PDF + Markdown).

The brief asks for ONE document containing the video link, the prompts, the
model name, an explanation of how the video was generated, and how the clips
were combined. The prompts and parameters are extracted from
Marwadi_2050_Video.ipynb rather than retyped, so the document always matches
what the notebook actually runs.

Run:  python make_document.py
"""

from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

import nbformat as nbf
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).parent
NB = HERE / "Marwadi_2050_Video.ipynb"

STUDENT = dict(name="Aditya Raj", enrol="92301733062", dept="ICT", batch="7EK1'A'")
VIDEO_LINK_PLACEHOLDER = "<<< PASTE YOUR PUBLIC GOOGLE DRIVE VIDEO LINK HERE >>>"

HEADER_FILL, ZEBRA, NOTE_FILL = "1F3864", "F7F7F5", "FFF6E5"
MARGIN = 0.7
USABLE = 8.5 - 2 * MARGIN


# ---------------------------------------------------------------- extract
def load_notebook_facts() -> dict:
    nb = nbf.read(NB, as_version=4)
    src = {i: c.source for i, c in enumerate(nb.cells) if c.cell_type == "code"}

    def cell_with(needle):
        return next(s for s in src.values() if needle in s)

    ns: dict = {}
    import math
    ns["math"] = math
    for line in cell_with("TARGET_FRAMES").splitlines():
        s = line.strip()
        if re.match(r"^(FPS|DURATION_S|N_SHOTS|TARGET_FRAMES|XFADE|JOINS|FRAMES_CLIP|WIDTH, HEIGHT)\s*=", s):
            exec(s, ns)

    p = cell_with("STORYBOARD = [")
    exec(p[:p.index("for s in STORYBOARD")], ns)

    g = cell_with("SEEDS = [")
    for line in g.splitlines():
        s = line.strip()
        if re.match(r"^(SEEDS|STEPS|GUIDANCE)\s*=", s):
            exec(s, ns)

    models = re.findall(r'"([^"]+)",\s*#\s*(primary|fallback)', cell_with("CANDIDATES = ["))
    return {**ns, "MODELS": models}


# ------------------------------------------------------------------ docx
def shade(el, fill):
    s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear"); s.set(qn("w:fill"), fill); el.append(s)


def add_table(doc, rows, widths, font=8.5, header=True):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for cell, val in zip(cells, row):
            cell.text = ""; cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            par = cell.paragraphs[0]; par.paragraph_format.space_after = Pt(0)
            run = par.add_run(str(val)); run.font.size = Pt(font)
            if i == 0 and header:
                run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade(cell._tc.get_or_add_tcPr(), HEADER_FILL)
            elif i % 2 == 0:
                shade(cell._tc.get_or_add_tcPr(), ZEBRA)
    pr = t._tbl.tblPr
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); pr.append(lay)
    w = OxmlElement("w:tblW"); w.set(qn("w:w"), str(int(USABLE * 1440))); w.set(qn("w:type"), "dxa"); pr.append(w)
    scale = USABLE / sum(widths)
    grid = t._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for col, ww in zip(grid.findall(qn("w:gridCol")), widths):
            col.set(qn("w:w"), str(int(ww * scale * 1440)))
    t.autofit = False
    for row in t.rows:
        for cell, ww in zip(row.cells, widths):
            cell.width = Inches(ww * scale)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def para(doc, text, size=10, bold=False, italic=False, justify=True):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for chunk, b in re.findall(r"(\*\*.+?\*\*)|([^*]+)", text):
        if chunk:
            r = p.add_run(chunk.strip("*")); r.bold = True
        elif b:
            r = p.add_run(b)
        else:
            continue
        r.font.size = Pt(size); r.italic = italic
        if bold:
            r.bold = True
    return p


def code_para(doc, text, size=8):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.font.name = "Consolas"; r.font.size = Pt(size)
    shade(p._p.get_or_add_pPr(), "F4F4F2")


def build(facts: dict) -> Path:
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(MARGIN)
        s.top_margin = s.bottom_margin = Inches(0.6)
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10)
    n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.12
    for name, sz in (("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 11)):
        st = doc.styles[name]; st.font.name = "Calibri"; st.font.size = Pt(sz)
        st.font.bold = True; st.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(4)

    doc.add_heading("15-Second AI Video Challenge", level=1)
    para(doc, "**Theme: Marwadi University in 2050**", size=12, justify=False)
    add_table(doc, [["Field", "Details"],
                    ["Student Name", STUDENT["name"]],
                    ["Enrollment No.", STUDENT["enrol"]],
                    ["Department", STUDENT["dept"]],
                    ["Batch", STUDENT["batch"]],
                    ["Course", "Prompt Engineering for Generative AI"]],
              widths=[1.6, 5.5], font=9)

    # 1. video link
    doc.add_heading("1. Final Video Link", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(VIDEO_LINK_PLACEHOLDER)
    r.font.size = Pt(11); r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    shade(p._p.get_or_add_pPr(), NOTE_FILL)
    para(doc, "Duration 15.00 seconds · {}×{} · {} fps · {} frames. Upload "
              "marwadi_2050_15s.mp4 to Google Drive, set sharing to \u201cAnyone with "
              "the link\u201d, and replace the line above before exporting this "
              "document to PDF.".format(facts["WIDTH"], facts["HEIGHT"],
                                        facts["FPS"], facts["TARGET_FRAMES"]),
         size=9, italic=True)

    # 2. model
    doc.add_heading("2. Video Generation Model", level=2)
    primary = facts["MODELS"][0][0] if facts["MODELS"] else "cerspense/zeroscope_v2_576w"
    add_table(doc, [["Item", "Details"],
                    ["Model", primary],
                    ["Type", "Open-source text-to-video latent diffusion"],
                    ["Base", "Fine-tuned from ModelScope text-to-video"],
                    ["Why chosen", "Watermark-free, clean 576×320 output, and fits a free "
                                   "Colab T4. A larger model such as CogVideoX gives better "
                                   "fidelity but risks running out of memory mid-render."],
                    ["Fallback", facts["MODELS"][1][0] if len(facts["MODELS"]) > 1
                                 else "damo-vilab/text-to-video-ms-1.7b"],
                    ["Framework", "Hugging Face diffusers (DiffusionPipeline)"],
                    ["Hardware", "Google Colab, NVIDIA T4"]],
              widths=[1.5, 5.6], font=9)

    # 3. structure
    doc.add_heading("3. Structure of the Video", level=2)
    para(doc, "The video uses the **5 × 3-second clips** option from the brief. "
              "Five clips of {} frames each are generated at {} fps and combined "
              "into one continuous file of exactly {} frames — 15.00 seconds."
              .format(facts["FRAMES_CLIP"], facts["FPS"], facts["TARGET_FRAMES"]))
    para(doc, "The five shots form a single narrative arc — **arrive → enter → "
              "learn → build → graduate** — so the sequence reads as a story "
              "rather than five unrelated clips.")
    rows = [["Shot", "Time", "Beat", "Camera movement"]]
    for s in facts["STORYBOARD"]:
        rows.append([s["n"], f"{(s['n']-1)*3}–{s['n']*3} s", s["beat"], s["camera"]])
    add_table(doc, rows, widths=[0.5, 0.8, 1.1, 4.7], font=9)

    # 4. prompts
    doc.add_heading("4. Prompts Used", level=2)
    para(doc, "Every prompt is composed from **six explicit control dimensions** — "
              "subject, environment, camera, lighting, style and quality — so each "
              "shot is directed rather than merely described.")
    add_table(doc, [["Dimension", "Purpose"],
                    ["Subject", "Who or what is in frame"],
                    ["Environment", "Where, and what the world looks like"],
                    ["Camera", "Shot type and its movement"],
                    ["Lighting", "Time of day and mood"],
                    ["Style", "Visual treatment"],
                    ["Quality", "Technical tags"]],
              widths=[1.3, 5.8], font=9)

    para(doc, "**Continuity anchor** — repeated verbatim in all five prompts so the "
              "five clips read as one place and one film:")
    code_para(doc, facts["CONTINUITY"])
    para(doc, "**Style and quality tags** — appended to every prompt:")
    code_para(doc, facts["STYLE"] + ",\n" + facts["QUALITY"])
    para(doc, "**Negative prompt** — applied to every shot:")
    code_para(doc, facts["NEGATIVE"])

    doc.add_heading("Full prompt for each shot", level=3)
    for s in facts["STORYBOARD"]:
        para(doc, f"**Shot {s['n']} — {s['beat']} ({(s['n']-1)*3}–{s['n']*3} s)**",
             size=10, justify=False)
        code_para(doc, facts["build_prompt"](s))

    # 5. how generated
    doc.add_heading("5. How the Video Was Generated", level=2)
    para(doc, "1. **Model load.** Zeroscope v2 576w is loaded through diffusers with "
              "fp16 weights, VAE slicing and model CPU offload so it fits a free T4.")
    para(doc, "2. **Per-shot generation.** Each of the five prompts is rendered "
              "independently to {} frames at {}×{}, using {} inference steps and "
              "guidance scale {}. Each shot has its own fixed seed ({}), so the "
              "whole render is reproducible."
              .format(facts["FRAMES_CLIP"], facts["WIDTH"], facts["HEIGHT"],
                      facts["STEPS"], facts["GUIDANCE"],
                      ", ".join(map(str, facts["SEEDS"]))))
    para(doc, "3. **Individual clips saved.** Every shot is written to its own MP4 "
              "in clips/ before anything is joined — those files are the evidence "
              "that five separate clips were generated.")
    para(doc, "4. **Combination.** The clips are concatenated into one continuous "
              "MP4 (section 6).")
    para(doc, "5. **Verification.** The notebook asserts the final duration is "
              "within the 15-second requirement before finishing.")
    add_table(doc, [["Parameter", "Value"],
                    ["Frames per clip", facts["FRAMES_CLIP"]],
                    ["Frame rate", f"{facts['FPS']} fps"],
                    ["Resolution", f"{facts['WIDTH']}×{facts['HEIGHT']}"],
                    ["Inference steps", facts["STEPS"]],
                    ["Guidance scale", facts["GUIDANCE"]],
                    ["Seeds", ", ".join(map(str, facts["SEEDS"]))],
                    ["Cross-dissolve", f"{facts['XFADE']} frames per join"],
                    ["Final length", f"{facts['TARGET_FRAMES']} frames = "
                                     f"{facts['TARGET_FRAMES']/facts['FPS']:.2f} s"]],
              widths=[2.0, 5.1], font=9)

    # 6. combination
    doc.add_heading("6. How the Clips Were Combined", level=2)
    para(doc, "The five clips are joined by concatenating their frame arrays and "
              "writing a single MP4 with imageio — no external editor was used, so "
              "the combination step is fully reproducible from the notebook.")
    para(doc, "**The frame arithmetic is the part worth explaining.** A "
              "cross-dissolve *blends* frames from both sides of a join rather than "
              "adding new ones, so it shortens the result. Five 24-frame clips with "
              "a 6-frame dissolve at each of the four joins would give "
              "5×24 − 4×6 = 96 frames = 12 seconds, not 15. Two measures prevent that:")
    para(doc, "• Frames per clip is derived from the target — ceil((120 + joins × "
              "XFADE) / 5) — so enough footage is generated to absorb the dissolves.")
    para(doc, "• The concatenation is then trimmed to exactly 120 frames.")
    para(doc, "The submitted version uses {} so the five shots meet at hard cuts, "
              "which is the normal convention for a short montage and lands on "
              "exactly 120 frames. Raising XFADE produces dissolves instead; the "
              "duration stays 15 seconds either way."
              .format("XFADE = 0" if facts["XFADE"] == 0
                      else f"XFADE = {facts['XFADE']}"))
    code_para(doc,
              "clip 1  [{f} frames]  0.0-3.0s\n"
              "clip 2  [{f} frames]  3.0-6.0s\n"
              "clip 3  [{f} frames]  6.0-9.0s\n"
              "clip 4  [{f} frames]  9.0-12.0s\n"
              "clip 5  [{f} frames]  12.0-15.0s\n"
              "                |\n"
              "     concatenate frame arrays ({x}-frame blend at each join)\n"
              "                |\n"
              "     trim to {t} frames  ->  marwadi_2050_15s.mp4  (15.00 s)"
              .format(f=facts["FRAMES_CLIP"], x=facts["XFADE"], t=facts["TARGET_FRAMES"]))

    # 7. prompt engineering
    doc.add_heading("7. Prompt Engineering — What Actually Controlled the Output", level=2)
    para(doc, "The brief states the objective is not a beautiful video but "
              "demonstrable control over scene, characters, environment, camera "
              "movement and visual storytelling. These were the changes that "
              "produced that control:")
    add_table(doc, [
        ["#", "Problem observed", "Change made", "Result"],
        ["1", "Clips looked like five unrelated places",
         "Added a shared continuity anchor to every prompt", "Reads as one campus"],
        ["2", "Almost no motion — nearly a still image",
         "Gave camera movement its own explicit clause", "Consistent directed motion"],
        ["3", "Warped faces in crowd shots",
         "Added deformed faces, extra limbs to the negative prompt", "Fewer artefacts"],
        ["4", "Colour jumped between shots",
         "Fixed the colour grade inside the continuity anchor", "Consistent grade"],
        ["5", "Text and watermarks appearing",
         "Added text, watermark, logo, subtitles to the negative prompt", "Clean frames"],
        ["6", "Weak prompt adherence",
         "Raised guidance scale to " + str(facts["GUIDANCE"]), "Shots match direction"],
    ], widths=[0.35, 2.1, 2.6, 2.05], font=8.5)
    para(doc, "**The most useful lesson:** in text-to-video, camera movement is a "
              "prompt component, not an afterthought. Describing a scene yields a "
              "nearly static shot; describing a camera moving through the scene is "
              "what produces motion. Storytelling here is carried by the camera — "
              "push-in, track, dolly, orbit, crane-up — not by cutting between "
              "unrelated images.")

    doc.add_heading("8. Files Submitted", level=2)
    add_table(doc, [["File", "Contents"],
                    ["marwadi_2050_15s.mp4", "The final 15-second video"],
                    ["clips/shot1..shot5.mp4", "The five individual 3-second clips"],
                    ["storyboard_contact_sheet.png", "Middle frame of each shot"],
                    ["video_manifest.json", "Every prompt, seed and parameter from the run"],
                    ["Marwadi_2050_Video.ipynb", "The notebook that generated everything"]],
              widths=[2.2, 4.9], font=9)

    OUT = HERE / "SUBMISSION_DOCUMENT.docx"
    doc.save(str(OUT))
    return OUT


def to_pdf(docx_path: Path) -> Path | None:
    out = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(out))
        if out.exists():
            return out
    except Exception:
        pass
    for exe in (r"C:\Program Files\LibreOffice\program\soffice.exe", "soffice"):
        try:
            subprocess.run([exe, "--headless", "--convert-to", "pdf",
                            "--outdir", str(out.parent), str(docx_path)],
                           check=True, capture_output=True, timeout=300)
            if out.exists():
                return out
        except (FileNotFoundError, subprocess.SubprocessError):
            continue
    return None


def main() -> int:
    facts = load_notebook_facts()
    print(f"Extracted from notebook: {len(facts['STORYBOARD'])} shots, "
          f"{facts['FRAMES_CLIP']} frames/clip, {facts['TARGET_FRAMES']} target frames")
    docx = build(facts)
    print(f"Wrote {docx.name}")
    pdf = to_pdf(docx)
    print(f"Wrote {pdf.name}" if pdf else "PDF skipped (no converter found)")
    print(f"\nRemember to replace the video link placeholder before submitting.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
