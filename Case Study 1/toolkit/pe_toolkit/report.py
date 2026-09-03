"""Markdown and DOCX renderers, plus PDF conversion."""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from . import config, document
from .blocks import (
    Bullets,
    Callout,
    Code,
    Doc,
    Figure,
    Heading,
    KeyValue,
    PageBreak,
    Para,
    Table,
    split_bold,
)

HEADER_FILL = "1F3864"
CALLOUT_FILL = "EEF3FA"
CODE_FILL = "F4F4F2"
ZEBRA_FILL = "F7F7F5"
RULE_COLOUR = "BFBFBF"

# Letter (8.5in) minus the left and right margins set in render_docx.
PAGE_WIDTH_IN = 8.5
MARGIN_IN = 0.65
USABLE_WIDTH_IN = PAGE_WIDTH_IN - 2 * MARGIN_IN  # 7.2in

ZWSP = "​"


def _wrappable(text: str, limit: int = 22) -> str:
    """Give Word somewhere to break long unbroken tokens.

    URLs and file paths contain no spaces, so Word refuses to wrap them and
    the containing cell pushes the whole table past the page margin. Inserting
    zero-width spaces after the natural separators fixes the overflow without
    changing the visible text or breaking copy-paste.
    """
    out = []
    for token in text.split(" "):
        if len(token) > limit:
            for separator in ("/", "?", "&", "=", "_", "-", "."):
                token = token.replace(separator, separator + ZWSP)
            # Still-oversized runs (long hashes) get a hard break every `limit`.
            pieces = token.split(ZWSP)
            token = ZWSP.join(
                ZWSP.join(p[i : i + limit] for i in range(0, len(p), limit))
                if len(p) > limit
                else p
                for p in pieces
            )
        out.append(token)
    return " ".join(out)


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


def _md_cell(value: object) -> str:
    text = "" if value is None else str(value)
    return text.replace("|", "\\|").replace("\n", " ").strip()


def _md_table(frame: pd.DataFrame) -> str:
    header = "| " + " | ".join(_md_cell(c) for c in frame.columns) + " |"
    rule = "| " + " | ".join("---" for _ in frame.columns) + " |"
    rows = [
        "| " + " | ".join(_md_cell(v) for v in row) + " |"
        for row in frame.itertuples(index=False)
    ]
    return "\n".join([header, rule, *rows])


def render_markdown(doc: Doc, out: Path) -> Path:
    lines: list[str] = []

    for block in doc.blocks:
        match block:
            case Heading(level=level, text=text):
                lines += ["", f"{'#' * level} {text}", ""]
            case Para(text=text):
                lines += [text, ""]
            case Bullets(items=items):
                lines += [f"- {i}" for i in items] + [""]
            case Table(frame=frame, caption=caption):
                if caption:
                    lines += [f"*{caption}*", ""]
                lines += [_md_table(frame), ""]
            case KeyValue(pairs=pairs, headers=headers, caption=caption):
                if caption:
                    lines += [f"*{caption}*", ""]
                lines += [
                    _md_table(pd.DataFrame(list(pairs), columns=list(headers))),
                    "",
                ]
            case Figure(path=path, caption=caption):
                rel = Path("figures") / path.name
                lines += [f"![{caption}]({rel.as_posix()})", "", f"*{caption}*", ""]
            case Code(text=text, caption=caption):
                if caption:
                    lines += [f"*{caption}*", ""]
                lines += ["```text", text, "```", ""]
            case Callout(title=title, text=text):
                lines += [f"> **{title}**", ">", f"> {text}", ""]
            case PageBreak():
                lines += ["---", ""]

    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    return out


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _shade(element, fill: str) -> None:
    pr = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def _shade_cell(cell, fill: str) -> None:
    _shade(cell._tc.get_or_add_tcPr(), fill)


def _shade_paragraph(paragraph, fill: str) -> None:
    _shade(paragraph._p.get_or_add_pPr(), fill)


def _border_paragraph(paragraph, colour: str) -> None:
    """Hairline box around a shaded paragraph so the fill has a defined edge."""
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "6")
        edge.set(qn("w:color"), colour)
        borders.append(edge)
    paragraph._p.get_or_add_pPr().append(borders)


def _repeat_header(row) -> None:
    """Mark a table row as a header so it repeats across page breaks."""
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _keep_together(row) -> None:
    """Stop a row splitting mid-cell across a page break."""
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _fix_layout(table, widths) -> None:
    """Pin the table to the text width and stop Word re-flowing the columns.

    `table.autofit = False` alone is not enough - Word still recomputes column
    widths from content unless the table declares a fixed layout and an
    explicit total width.
    """
    tbl_pr = table._tbl.tblPr

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    total = OxmlElement("w:tblW")
    total.set(qn("w:w"), str(int(USABLE_WIDTH_IN * 1440)))
    total.set(qn("w:type"), "dxa")
    tbl_pr.append(total)

    margins = OxmlElement("w:tblCellMar")
    for side, value in (("top", 40), ("left", 70), ("bottom", 40), ("right", 70)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)

    if not widths:
        widths = [1] * len(table.columns)

    table.autofit = False
    scale = USABLE_WIDTH_IN / sum(widths)
    inches = [w * scale for w in widths]

    # Both the grid and every cell need the width, or Word ignores it.
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for column, width in zip(grid.findall(qn("w:gridCol")), inches):
            column.set(qn("w:w"), str(int(width * 1440)))

    for row in table.rows:
        for cell, width in zip(row.cells, inches):
            cell.width = Inches(width)


def _add_table(docx, frame: pd.DataFrame, widths=None, font_pt: float | None = None) -> None:
    if frame.empty:
        return

    # Dense tables get a smaller face so the columns stay legible.
    if font_pt is None:
        font_pt = 7.5 if len(frame.columns) >= 5 else 8.5

    table = docx.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0]
    for cell, name in zip(header.cells, frame.columns):
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(_wrappable(str(name)))
        run.bold = True
        run.font.size = Pt(font_pt)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, HEADER_FILL)
    _repeat_header(header)

    for i, values in enumerate(frame.itertuples(index=False)):
        row = table.add_row()
        _keep_together(row)
        cells = row.cells
        for cell, value in zip(cells, values):
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(
                "" if value is None else _wrappable(str(value))
            )
            run.font.size = Pt(font_pt)
            if i % 2 == 1:
                _shade_cell(cell, ZEBRA_FILL)

    _fix_layout(table, widths)

    spacer = docx.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def _add_rich_paragraph(docx, text: str, *, size: float = 10.5) -> None:
    """Add prose, one Word paragraph per blank-line-separated block.

    The narrative in `content.py` is written as triple-quoted strings wrapped
    at column ~76. Those source line breaks are not meaningful, but python-docx
    turns a `\\n` inside a run into a hard line break, and justified text then
    stretches every one of those short lines across the full column. Collapsing
    intra-paragraph whitespace first is what stops that.
    """
    for block in re.split(r"\n\s*\n", text.strip()):
        collapsed = re.sub(r"\s+", " ", block).strip()
        if not collapsed:
            continue
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for chunk, bold in split_bold(collapsed):
            run = paragraph.add_run(chunk)
            run.bold = bold
            run.font.size = Pt(size)


def _add_page_numbers(section) -> None:
    """Footer of the form 'Page N'. Field codes, so Word keeps them live."""
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x89, 0x87, 0x81)

    field = paragraph.add_run()
    field.font.size = Pt(8.5)
    field.font.color.rgb = RGBColor(0x89, 0x87, 0x81)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, end):
        field._r.append(node)


def _style_headings(docx) -> None:
    """Readable, tightly spaced headings instead of Word's airy defaults."""
    sizes = {
        "Heading 1": (18, True, 0, 10),
        "Heading 2": (14, True, 14, 6),
        "Heading 3": (11.5, True, 10, 4),
        "Heading 4": (10, True, 8, 3),
    }
    for name, (size, bold, before, after) in sizes.items():
        style = docx.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def render_docx(doc: Doc, out: Path) -> Path:
    docx = DocxDocument()

    for section in docx.sections:
        section.left_margin = section.right_margin = Inches(MARGIN_IN)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        _add_page_numbers(section)

    normal = docx.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    _style_headings(docx)

    # Justified text without hyphenation opens rivers of white space on the
    # long technical words this report is full of.
    hyphenation = OxmlElement("w:autoHyphenation")
    hyphenation.set(qn("w:val"), "true")
    docx.settings.element.append(hyphenation)

    for block in doc.blocks:
        match block:
            case Heading(level=level, text=text):
                docx.add_heading(text, level=min(level, 4))
            case Para(text=text):
                _add_rich_paragraph(docx, text)
            case Bullets(items=items):
                for item in items:
                    docx.add_paragraph(item, style="List Bullet")
            case Table(frame=frame, caption=caption, widths=widths):
                if caption:
                    _add_rich_paragraph(docx, f"**{caption}**", size=9)
                _add_table(docx, frame, widths)
            case KeyValue(pairs=pairs, headers=headers, caption=caption):
                if caption:
                    _add_rich_paragraph(docx, f"**{caption}**", size=9)
                _add_table(
                    docx,
                    pd.DataFrame(list(pairs), columns=list(headers)),
                    widths=[2.2, 6.4],
                )
            case Figure(path=path, caption=caption, width_in=width_in):
                if path.exists():
                    docx.add_picture(str(path), width=Inches(width_in))
                    picture = docx.paragraphs[-1]
                    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Keep the figure with its caption.
                    picture.paragraph_format.keep_with_next = True
                    picture.paragraph_format.space_before = Pt(6)
                caption_paragraph = docx.add_paragraph()
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption_paragraph.add_run(caption)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x52, 0x51, 0x4E)
            case Code(text=text, caption=caption):
                if caption:
                    _add_rich_paragraph(docx, f"**{caption}**", size=9)
                paragraph = docx.add_paragraph()
                fmt = paragraph.paragraph_format
                fmt.space_before = Pt(3)
                fmt.space_after = Pt(9)
                fmt.left_indent = Inches(0.12)
                fmt.right_indent = Inches(0.12)
                fmt.line_spacing = 1.0
                run = paragraph.add_run(text)
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                _shade_paragraph(paragraph, CODE_FILL)
                _border_paragraph(paragraph, RULE_COLOUR)
            case Callout(title=title, text=text):
                paragraph = docx.add_paragraph()
                fmt = paragraph.paragraph_format
                fmt.space_before = Pt(6)
                fmt.space_after = Pt(10)
                fmt.left_indent = Inches(0.1)
                fmt.right_indent = Inches(0.1)
                title_run = paragraph.add_run(f"{title}\n")
                title_run.bold = True
                title_run.font.size = Pt(9.5)
                body_run = paragraph.add_run(text)
                body_run.font.size = Pt(9.5)
                _shade_paragraph(paragraph, CALLOUT_FILL)
                _border_paragraph(paragraph, "9DB2CE")
            case PageBreak():
                docx.add_page_break()

    out.parent.mkdir(parents=True, exist_ok=True)
    docx.save(str(out))
    return out


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def render_pdf(docx_path: Path, out: Path) -> Path | None:
    """Convert the DOCX to PDF using whichever converter is present.

    Word is preferred because it preserves the table styling; LibreOffice is
    the fallback. Returns None if neither is available, rather than failing the
    whole build.
    """
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(out))
        if out.exists():
            return out
    except Exception:
        pass

    for candidate in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        "soffice",
        "libreoffice",
    ):
        try:
            subprocess.run(
                [candidate, "--headless", "--convert-to", "pdf",
                 "--outdir", str(out.parent), str(docx_path)],
                check=True,
                capture_output=True,
                timeout=300,
            )
            produced = out.parent / (docx_path.stem + ".pdf")
            if produced.exists():
                if produced != out:
                    produced.replace(out)
                return out
        except (FileNotFoundError, subprocess.SubprocessError):
            continue

    return None


# ---------------------------------------------------------------------------


def build_all(directory: Path | None = None) -> dict[str, Path | None]:
    directory = directory or config.REPORT_DIR
    directory.mkdir(parents=True, exist_ok=True)

    doc = document.build()
    stem = "CASE_STUDY_REPORT"

    markdown = render_markdown(doc, directory / f"{stem}.md")
    docx_path = render_docx(doc, directory / f"{stem}.docx")
    pdf = render_pdf(docx_path, directory / f"{stem}.pdf")

    return {"markdown": markdown, "docx": docx_path, "pdf": pdf}
