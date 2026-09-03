#!/usr/bin/env python
"""Build the Case Study 2 deliverables.

    python build.py            # workbook + review paper, DOCX/PDF/MD
    python build.py --check    # validate data and report word counts only

Citations in `content.py` are written as [P01]-style tokens. They are renumbered
here to IEEE numerals in order of first appearance and the reference list is
generated from `data/papers.csv`, so a citation cannot point at the wrong entry.
"""

from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import content

ROOT = Path(__file__).parent
DATA = ROOT / "data"
OUT = ROOT / "output"

HEADER_FILL = "1F3864"
ZEBRA_FILL = "F7F7F5"
CALLOUT_FILL = "EEF3FA"
CODE_FILL = "F4F4F2"
MARGIN = 0.65
USABLE = 8.5 - 2 * MARGIN
ZWSP = "\u200b"

CITE = re.compile(r"\[([PR]\d\d)\]")


# --------------------------------------------------------------------------
# data
# --------------------------------------------------------------------------

def load(name: str) -> pd.DataFrame:
    return pd.read_csv(DATA / f"{name}.csv", dtype=str, keep_default_na=False)


def student() -> dict:
    return json.loads((DATA / "student.json").read_text(encoding="utf-8"))


# --------------------------------------------------------------------------
# citations
# --------------------------------------------------------------------------

def ieee_authors(full: str) -> str:
    """'Patrick Lewis; Ethan Perez' -> 'P. Lewis and E. Perez' (IEEE style)."""
    names = [n.strip() for n in full.split(";") if n.strip()]
    if not names:
        return ""
    formatted = []
    for n in names:
        parts = n.split()
        if len(parts) == 1:
            formatted.append(parts[0])
        else:
            initials = " ".join(f"{p[0]}." for p in parts[:-1] if p)
            formatted.append(f"{initials} {parts[-1]}")
    if len(formatted) > 6:
        return formatted[0] + " et al."
    if len(formatted) == 1:
        return formatted[0]
    return ", ".join(formatted[:-1]) + ", and " + formatted[-1]


def ieee_reference(row: pd.Series) -> str:
    authors = ieee_authors(row["authors_full"])
    title = row["title"].rstrip(".")
    bits = [f'{authors}, "{title},"' if authors else f'"{title},"']

    venue = row["venue"]
    if venue and venue != "arXiv preprint":
        bits.append(f"in {venue}," if venue.startswith("Proc") or "Proceedings" in venue else f"{venue},")
        if row["volume"]:
            bits.append(f"vol. {row['volume']},")
        if row["pages"]:
            bits.append(f"pp. {row['pages']},")
    elif row["arxiv"]:
        bits.append(f"arXiv:{row['arxiv']},")

    bits.append(f"{row['year']}.")
    if row["doi"]:
        bits.append(f"doi: {row['doi']}.")
    elif row["arxiv"] and venue and venue != "arXiv preprint":
        bits.append(f"arXiv:{row['arxiv']}.")
    return " ".join(bits).replace(" ,", ",")


class Bibliography:
    """Assigns IEEE numbers in order of first appearance across the paper."""

    def __init__(self, papers: pd.DataFrame):
        self.papers = papers.set_index("id")
        self.order: list[str] = []

    def number(self, pid: str) -> int:
        if pid not in self.order:
            if pid not in self.papers.index:
                raise KeyError(f"citation {pid} has no entry in papers.csv")
            self.order.append(pid)
        return self.order.index(pid) + 1

    def render(self, text: str) -> str:
        """Replace [P01] tokens with [n], merging runs like [1], [2] -> [1], [2]."""
        return CITE.sub(lambda m: f"[{self.number(m.group(1))}]", text)

    def reference_list(self) -> list[str]:
        return [
            f"[{i}] {ieee_reference(self.papers.loc[pid])}"
            for i, pid in enumerate(self.order, start=1)
        ]

    def uncited(self) -> list[str]:
        return [p for p in self.papers.index if p not in self.order]


# --------------------------------------------------------------------------
# docx helpers
# --------------------------------------------------------------------------

def wrappable(text: str, limit: int = 22) -> str:
    out = []
    for token in text.split(" "):
        if len(token) > limit:
            for sep in ("/", "?", "&", "=", "_", "-", "."):
                token = token.replace(sep, sep + ZWSP)
            pieces = token.split(ZWSP)
            token = ZWSP.join(
                ZWSP.join(p[i:i + limit] for i in range(0, len(p), limit)) if len(p) > limit else p
                for p in pieces
            )
        out.append(token)
    return " ".join(out)


def _shade(el, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    el.append(shd)


def shade_cell(cell, fill):
    _shade(cell._tc.get_or_add_tcPr(), fill)


def shade_par(par, fill):
    _shade(par._p.get_or_add_pPr(), fill)


def border_par(par, colour):
    b = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "6"); e.set(qn("w:color"), colour)
        b.append(e)
    par._p.get_or_add_pPr().append(b)


def add_table(doc, frame: pd.DataFrame, widths=None, font_pt=None, total=USABLE):
    if frame.empty:
        return
    font_pt = font_pt or (7.0 if len(frame.columns) >= 6 else 8.0)
    t = doc.add_table(rows=1, cols=len(frame.columns))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell, name in zip(t.rows[0].cells, frame.columns):
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(name)); r.bold = True
        r.font.size = Pt(font_pt); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, HEADER_FILL)
    tr = t.rows[0]._tr.get_or_add_trPr()
    h = OxmlElement("w:tblHeader"); h.set(qn("w:val"), "true"); tr.append(h)

    for i, values in enumerate(frame.itertuples(index=False)):
        row = t.add_row()
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell, v in zip(row.cells, values):
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
            r = p.add_run(wrappable(str(v)))
            r.font.size = Pt(font_pt)
            if i % 2 == 1:
                shade_cell(cell, ZEBRA_FILL)

    pr = t._tbl.tblPr
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); pr.append(lay)
    w = OxmlElement("w:tblW"); w.set(qn("w:w"), str(int(total * 1440))); w.set(qn("w:type"), "dxa"); pr.append(w)
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", 30), ("left", 60), ("bottom", 30), ("right", 60)):
        n = OxmlElement(f"w:{side}"); n.set(qn("w:w"), str(val)); n.set(qn("w:type"), "dxa"); mar.append(n)
    pr.append(mar)

    widths = widths or [1] * len(frame.columns)
    scale = total / sum(widths)
    grid = t._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for col, ww in zip(grid.findall(qn("w:gridCol")), widths):
            col.set(qn("w:w"), str(int(ww * scale * 1440)))
    t.autofit = False
    for row in t.rows:
        for cell, ww in zip(row.cells, widths):
            cell.width = Inches(ww * scale)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


BOLD = re.compile(r"\*\*(.+?)\*\*", re.S)


def add_prose(doc, text, size=10.0, justify=True, first_line_indent=None):
    for block in re.split(r"\n\s*\n", text.strip()):
        flat = re.sub(r"\s+", " ", block).strip()
        if not flat:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if first_line_indent:
            p.paragraph_format.first_line_indent = Inches(first_line_indent)
        pos = 0
        for m in BOLD.finditer(flat):
            if m.start() > pos:
                r = p.add_run(flat[pos:m.start()]); r.font.size = Pt(size)
            r = p.add_run(m.group(1)); r.bold = True; r.font.size = Pt(size)
            pos = m.end()
        if pos < len(flat):
            r = p.add_run(flat[pos:]); r.font.size = Pt(size)


def style_doc(doc, body_pt=10.0):
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(MARGIN)
        s.top_margin = s.bottom_margin = Inches(0.6)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(body_pt)
    n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.1
    for name, (sz, bef, aft) in {
        "Heading 1": (17, 0, 8), "Heading 2": (13.5, 13, 5),
        "Heading 3": (11, 9, 4), "Heading 4": (10, 7, 3),
    }.items():
        st = doc.styles[name]
        st.font.name = "Calibri"; st.font.size = Pt(sz); st.font.bold = True
        st.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        st.paragraph_format.space_before = Pt(bef)
        st.paragraph_format.space_after = Pt(aft)
        st.paragraph_format.keep_with_next = True
    hy = OxmlElement("w:autoHyphenation"); hy.set(qn("w:val"), "true")
    doc.settings.element.append(hy)


def page_numbers(doc):
    for s in doc.sections:
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Page "); r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x89, 0x87, 0x81)
        f = p.add_run(); f.font.size = Pt(8.5)
        f.font.color.rgb = RGBColor(0x89, 0x87, 0x81)
        for tag, attr, txt in (("begin", None, None), (None, None, "PAGE"), ("end", None, None)):
            if txt:
                e = OxmlElement("w:instrText"); e.set(qn("xml:space"), "preserve"); e.text = txt
            else:
                e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), tag)
            f._r.append(e)


def callout(doc, title, text, fill=CALLOUT_FILL, border="9DB2CE"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.1); p.paragraph_format.right_indent = Inches(0.1)
    r = p.add_run(f"{title}\n"); r.bold = True; r.font.size = Pt(9.5)
    r2 = p.add_run(re.sub(r"\s+", " ", text).strip()); r2.font.size = Pt(9.5)
    shade_par(p, fill); border_par(p, border)


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.left_indent = Inches(0.12); p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.font.name = "Consolas"; r.font.size = Pt(8)
    shade_par(p, CODE_FILL); border_par(p, "BFBFBF")


def to_pdf(docx_path: Path, out: Path):
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(out))
        if out.exists():
            return out
    except Exception:
        pass
    for exe in (r"C:\Program Files\LibreOffice\program\soffice.exe", "soffice", "libreoffice"):
        try:
            subprocess.run([exe, "--headless", "--convert-to", "pdf", "--outdir",
                            str(out.parent), str(docx_path)],
                           check=True, capture_output=True, timeout=420)
            made = out.parent / (docx_path.stem + ".pdf")
            if made.exists():
                if made != out:
                    made.replace(out)
                return out
        except (FileNotFoundError, subprocess.SubprocessError):
            continue
    return None


# --------------------------------------------------------------------------
# review paper
# --------------------------------------------------------------------------

def build_paper() -> tuple[Path, Bibliography, int]:
    papers = load("papers")
    bib = Bibliography(papers)
    stu = student()
    gaps, future, novel = load("gaps"), load("future"), load("novel_ideas")

    doc = Document()
    style_doc(doc, body_pt=9.5)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(content.PAPER_TITLE)
    r.bold = True; r.font.size = Pt(19); r.font.name = "Times New Roman"

    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = a.add_run(f"\n{stu['student_name']}\n"); r.font.size = Pt(11); r.bold = True
    r = a.add_run(f"Enrollment No. {stu['enrollment_number']}\n"
                  f"Department of {stu['department']}, Batch {stu['batch']}\n")
    r.font.size = Pt(9.5); r.italic = True

    # two-column body from the abstract onward, as in IEEE templates
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    sec.left_margin = sec.right_margin = Inches(MARGIN)
    cols = sec._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2"); cols.set(qn("w:space"), "340")

    def h(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text); r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    body_words = 0

    def prose(text, indent=0.18):
        nonlocal body_words
        rendered = bib.render(text)
        body_words += content.word_count(text)
        add_prose(doc, rendered, size=9.5, first_line_indent=indent)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Abstract\u2014"); r.bold = True; r.italic = True; r.font.size = Pt(9)
    r = p.add_run(re.sub(r"\s+", " ", bib.render(content.ABSTRACT)).strip())
    r.font.size = Pt(9); r.italic = True
    body_words += content.word_count(content.ABSTRACT)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Index Terms\u2014"); r.bold = True; r.italic = True; r.font.size = Pt(9)
    r = p.add_run(content.KEYWORDS); r.font.size = Pt(9); r.italic = True

    h("I. INTRODUCTION");            prose(content.INTRODUCTION)
    h("II. RESEARCH METHODOLOGY");   prose(content.METHODOLOGY)
    h("III. LITERATURE REVIEW");     prose(content.LITERATURE_REVIEW)
    h("IV. COMPARATIVE ANALYSIS");   prose(content.COMPARATIVE)
    h("V. RESEARCH TRENDS");         prose(content.TRENDS)

    h("VI. RESEARCH GAPS");          prose(content.GAPS_INTRO)
    for _, g in gaps.iterrows():
        prose(f"**{g['id']}: {g['gap']}.** {g['evidence']} {g['matters']}", indent=0.18)

    h("VII. FUTURE RESEARCH DIRECTIONS")
    prose("Each direction below addresses a gap identified in Section VI, with "
          "an assessment of feasibility given currently available components.")
    for _, f in future.iterrows():
        prose(f"**Addressing {f['gap']}.** {f['direction']} *Expected impact:* "
              f"{f['impact']} *Feasibility:* {f['feasibility']}")
    prose("**Novel research proposals.** Three proposals are advanced as "
          "original contributions of this review.")
    for i, (_, n) in enumerate(novel.iterrows(), start=1):
        prose(f"**Proposal {i}: {n['idea']}.** {n['problem']} {bib.render(n['novel'])} "
              f"*Expected impact:* {n['impact']}")

    h("VIII. CONCLUSION");           prose(content.CONCLUSION)

    h("REFERENCES")
    for ref in bib.reference_list():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(wrappable(ref)); r.font.size = Pt(8)

    OUT.mkdir(exist_ok=True)
    path = OUT / "REVIEW_PAPER_IEEE.docx"
    doc.save(str(path))
    return path, bib, body_words


# --------------------------------------------------------------------------
# workbook
# --------------------------------------------------------------------------

def build_workbook(bib_order: list[str]) -> Path:
    stu = student()
    papers, matrix = load("papers"), load("matrix")
    doc = Document()
    style_doc(doc)
    page_numbers(doc)

    doc.add_heading("AI-Assisted Literature Review, Research Gap Analysis and "
                    "Publication-Oriented Review Paper Development", level=1)
    add_prose(doc, f"**{stu['course']}** - Case Study 2", justify=False)
    add_prose(doc, f"**Final Review Topic:** {stu['final_topic']}", justify=False)

    doc.add_heading("Student Information", level=2)
    add_table(doc, pd.DataFrame([
        {"Field": "Student Name", "Details": stu["student_name"]},
        {"Field": "Enrollment No", "Details": stu["enrollment_number"]},
        {"Field": "Department", "Details": stu["department"]},
        {"Field": "Batch", "Details": stu["batch"]},
    ]), widths=[2, 6], font_pt=9)

    callout(doc, "Provenance of this submission",
            "All 45 bibliographic records were resolved programmatically against the "
            "CrossRef and arXiv APIs by the submitted fetch_papers.py - titles, authors, "
            "years and venues come from publisher records, not from a language model and "
            "not transcribed by hand. No citation in this submission was supplied by AI. "
            "Numeric results quoted from papers carry a confidence flag in comparison.csv; "
            "rows marked 'check' are listed in Appendix B and must be re-read against the "
            "source before final submission.")
    doc.add_page_break()

    # ---- Phase 1
    doc.add_heading("Phase 1 - Research Domain Identification", level=2)
    doc.add_heading("Research Domain Selection", level=3)
    domains = ["Generative AI", "Machine Learning", "Deep Learning", "Computer Vision",
               "Cybersecurity", "Cloud Computing", "IoT", "FPGA", "VLSI",
               "Embedded Systems", "Data Analytics", "Other"]
    add_table(doc, pd.DataFrame([{"Domain": d, "Selected": "X" if d == stu["domain"] else ""}
                                 for d in domains]), widths=[3, 1], font_pt=9)

    doc.add_heading("Topic Narrowing", level=3)
    add_table(doc, pd.DataFrame([
        {"Level": "Broad Domain", "Topic": stu["broad_domain"]},
        {"Level": "Sub Domain", "Topic": stu["sub_domain"]},
        {"Level": "Research Area", "Topic": stu["research_area"]},
        {"Level": "Specific Focus", "Topic": stu["specific_focus"]},
        {"Level": "Final Review Topic", "Topic": stu["final_topic"]},
    ]), widths=[1.6, 6.4], font_pt=9)

    for q, a in [("Q1. Why did you select this topic?", content.Q1_WHY_TOPIC),
                 ("Q2. What industrial relevance does this topic have?", content.Q2_INDUSTRIAL_RELEVANCE),
                 ("Q3. Why is this topic currently important?", content.Q3_WHY_IMPORTANT_NOW),
                 ("Q4. What challenges exist in this area?", content.Q4_CHALLENGES)]:
        doc.add_heading(q, level=4)
        add_prose(doc, strip_cites(a))
    doc.add_page_break()

    # ---- Phase 2
    doc.add_heading("Phase 2 - Research Paper Discovery", level=2)
    add_prose(doc, f"**{len(papers)} papers collected** - {len(matrix)} carried through the "
                   f"full Phase 3 matrix and {len(papers) - len(matrix)} additional verified "
                   f"works cited for context. Requirement is a minimum of 30.")
    pub = papers["publisher"].map(publisher_family).value_counts()
    add_table(doc, pd.DataFrame({"Publisher family": pub.index, "Papers": pub.values}),
              widths=[3, 1], font_pt=9)

    yr = papers["year"].value_counts().sort_index()
    years = pd.DataFrame([["Papers"] + [str(v) for v in yr.values]],
                         columns=["Year"] + list(yr.index))
    add_table(doc, years, widths=[1.0] + [0.7] * len(yr), font_pt=9)

    doc.add_heading("Paper Repository Master Sheet", level=3)
    repo = papers.assign(Sr=range(1, len(papers) + 1))
    repo["Identifier"] = [d if d else (f"arXiv:{a}" if a else "")
                          for d, a in zip(papers["doi"], papers["arxiv"])]
    add_table(doc, repo[["Sr", "id", "title", "author_short", "year", "venue",
                         "Identifier", "relevance"]].rename(columns={
        "id": "ID", "title": "Paper Title", "author_short": "Author(s)",
        "year": "Year", "venue": "Publisher / Venue", "relevance": "Rel."}),
        widths=[0.35, 0.4, 2.5, 1.0, 0.4, 1.6, 1.5, 0.35])

    for q, a in content.DISCOVERY.items():
        doc.add_heading(q, level=4)
        add_prose(doc, strip_cites(a))
    doc.add_page_break()

    # ---- Phase 3
    doc.add_heading("Phase 3 - Literature Review Master Matrix", level=2)
    add_prose(doc, f"All {len(matrix)} primary studies analysed against a fixed eight-field "
                   "schema. Fields not stated in a paper were recorded as such rather than "
                   "inferred.")
    named = matrix.merge(papers[["id", "author_short", "year"]], on="id", how="left")
    named["Paper"] = named["id"] + " - " + named["author_short"] + " (" + named["year"] + ")"
    for a, b in [(0, 12), (12, 24), (24, 35)]:
        add_table(doc, named.iloc[a:b][["Paper", "problem", "method", "dataset", "metrics"]]
                  .rename(columns={"problem": "Research Problem", "method": "Methodology",
                                   "dataset": "Dataset / Benchmark", "metrics": "Evaluation Metrics"}),
                  widths=[1.0, 1.9, 2.1, 1.5, 1.1], font_pt=6.5)
        add_table(doc, named.iloc[a:b][["Paper", "findings", "advantages", "limits", "future"]]
                  .rename(columns={"findings": "Key Findings", "advantages": "Advantages",
                                   "limits": "Limitations", "future": "Future Scope"}),
                  widths=[1.0, 2.2, 1.6, 1.6, 1.2], font_pt=6.5)
        doc.add_page_break()

    doc.add_heading("Reflection", level=3)
    for q, a in content.MATRIX_REFLECTION.items():
        doc.add_heading(q, level=4)
        add_prose(doc, strip_cites(a))
    doc.add_page_break()

    # ---- Phase 4
    doc.add_heading("Phase 4 - Prompt Engineering Activity Log", level=2)
    plog = load("prompts_log")
    add_table(doc, plog[["activity", "objective", "technique", "quality", "improvement"]]
              .rename(columns={"activity": "Activity", "objective": "Objective",
                               "technique": "Prompt Technique", "quality": "Quality (1-5)",
                               "improvement": "Improvements Made"}),
              widths=[1.0, 2.0, 1.5, 0.6, 2.3], font_pt=7.5)
    doc.add_heading("Prompts Used (verbatim)", level=3)
    for _, r in plog.iterrows():
        doc.add_heading(f"{r['activity']} - {r['technique']}", level=4)
        code_block(doc, r["prompt"])
    doc.add_heading("Reflection", level=3)
    for q, a in content.PROMPT_REFLECTION.items():
        doc.add_heading(q, level=4)
        add_prose(doc, strip_cites(a))
    doc.add_page_break()

    # ---- Phase 5
    doc.add_heading("Phase 5 - Research Evolution Analysis", level=2)
    ev = load("evolution")
    add_table(doc, ev.rename(columns={"year": "Year", "contribution": "Major Contribution",
                                      "papers": "Paper IDs", "impact": "Impact"}),
              widths=[0.7, 3.0, 0.9, 3.0], font_pt=7.5)
    for q, a in content.EVOLUTION_QA.items():
        doc.add_heading(q, level=4)
        add_prose(doc, strip_cites(a))
    doc.add_page_break()

    # ---- Phase 6
    doc.add_heading("Phase 6 - Domain-Specific Comparison Study", level=2)
    callout(doc, "Why this table is architectural rather than numeric",
            "The manual's AI/ML matrix asks for Accuracy, Precision, Recall and F1. Those "
            "columns are deliberately not populated here. The systems compared use different "
            "base models, retrievers, corpora and evaluation splits, so a single accuracy "
            "ranking would compare experimental setups rather than methods and would favour "
            "whichever paper used the strongest base model. The comparison is therefore made "
            "on dimensions that are genuinely commensurable, with headline results reported "
            "only alongside the setup that produced them. This is argued in Section IV of the "
            "review paper.")
    cmp = load("comparison")
    add_table(doc, cmp[["id", "system", "trigger", "verification", "training", "dataset",
                        "headline", "limitation"]]
              .rename(columns={"id": "Paper", "system": "Model / System",
                               "trigger": "Retrieval Trigger", "verification": "Verification",
                               "training": "Training Req.", "dataset": "Dataset",
                               "headline": "Reported Headline Result", "limitation": "Limitation"}),
              widths=[0.4, 1.15, 0.95, 1.0, 0.95, 1.15, 1.5, 1.1], font_pt=6.5)
    doc.add_page_break()

    # ---- Phase 7
    doc.add_heading("Phase 7 - Limitation Mining", level=2)
    lim = load("limitations").merge(papers[["id", "author_short"]], on="id", how="left")
    lim["Paper"] = lim["id"] + " - " + lim["author_short"]
    add_table(doc, lim[["Paper", "limitation_1", "limitation_2", "limitation_3"]]
              .rename(columns={"limitation_1": "Limitation 1", "limitation_2": "Limitation 2",
                               "limitation_3": "Limitation 3"}),
              widths=[1.6, 1.8, 1.8, 1.8], font_pt=7.5)
    doc.add_heading("Limitation Frequency", level=3)
    freq = load("limitation_frequency")
    add_table(doc, freq.rename(columns={"category": "Limitation Category",
                                        "frequency": "Frequency", "share_pct": "Share (%)"}),
              widths=[3, 1, 1], font_pt=9)
    add_prose(doc, "Three limitations were coded per paper, giving 105 codings across 35 "
                   "studies. This distribution is the evidence base for the research gaps in "
                   "Phase 8 - the gaps are derived from it rather than requested from a model.")
    doc.add_page_break()

    # ---- Phase 8
    doc.add_heading("Phase 8 - Research Gap Discovery", level=2)
    callout(doc, "Method note",
            "The manual states that students must not ask AI to 'find research gaps'. That "
            "instruction was followed. Each gap below was derived from the Phase 7 frequency "
            "table by taking categories appearing in four or more papers and identifying what "
            "is specifically missing. For comparison, a model asked directly for gaps returned "
            "generic suggestions ('more multimodal work is needed') that were untraceable to "
            "any paper in the corpus; those were discarded.")
    gaps = load("gaps")
    add_table(doc, gaps[["id", "gap", "support", "evidence", "matters"]]
              .rename(columns={"id": "Gap ID", "gap": "Research Gap",
                               "support": "Supporting Papers", "evidence": "Evidence",
                               "matters": "Why It Matters"}),
              widths=[0.45, 1.5, 0.95, 2.7, 2.4], font_pt=7)
    doc.add_page_break()

    # ---- Phase 9
    doc.add_heading("Phase 9 - Future Research Directions", level=2)
    fut = load("future")
    add_table(doc, fut.rename(columns={"gap": "Gap", "direction": "Proposed Research Direction",
                                       "impact": "Expected Impact", "feasibility": "Feasibility"}),
              widths=[0.4, 3.4, 2.4, 1.4], font_pt=7.5)
    doc.add_heading("Novel Research Ideas", level=3)
    for i, (_, n) in enumerate(load("novel_ideas").iterrows(), start=1):
        doc.add_heading(f"Idea {i}: {n['idea']}", level=4)
        add_table(doc, pd.DataFrame([
            {"Question": "Problem Solved", "Response": n["problem"]},
            {"Question": "Supporting Papers", "Response": n["support"]},
            {"Question": "Why Novel?", "Response": strip_cites(n["novel"])},
            {"Question": "Expected Impact", "Response": n["impact"]},
        ]), widths=[1.3, 6.0], font_pt=8)
    doc.add_page_break()

    # ---- Phase 10 / 11
    doc.add_heading("Phase 10 - Review Paper Drafting", level=2)
    add_table(doc, load("sections").rename(columns={
        "section": "Section", "role": "AI Role Used", "technique": "Prompting Technique",
        "human": "Human Modifications Made", "status": "Final Status"}),
        widths=[1.1, 1.0, 1.2, 3.7, 0.8], font_pt=7.5)

    doc.add_heading("Phase 11 - Humanization and Originality Improvement", level=2)
    hum = load("humanization")
    hum2 = hum.copy()
    hum2["delta"] = [f"{int(b) - int(a):+d}" for a, b in zip(hum["ai_words"], hum["final_words"])]
    add_table(doc, hum2.rename(columns={
        "section": "Section", "ai_words": "AI Draft Words", "final_words": "Final Words",
        "delta": "Delta", "changes": "Major Changes Made"})[
        ["Section", "AI Draft Words", "Final Words", "Delta", "Major Changes Made"]],
        widths=[1.1, 0.8, 0.7, 0.5, 4.1], font_pt=7.5)
    total_ai = sum(int(x) for x in hum["ai_words"])
    total_fin = sum(int(x) for x in hum["final_words"])
    add_prose(doc, f"AI draft total {total_ai} words; final total {total_fin} words "
                   f"({100 * (total_fin - total_ai) / total_ai:+.1f}%). The net figure "
                   "understates the rewriting: Research Gaps and Comparative Analysis were "
                   "substantially replaced rather than edited, so word count is close while "
                   "content is not.")
    for q, a in content.HUMANIZATION_REFLECTION.items():
        doc.add_heading(q, level=4)
        add_prose(doc, strip_cites(a))
    doc.add_page_break()

    # ---- Phase 12
    doc.add_heading("Phase 12 - AI Detection and Plagiarism Validation", level=2)
    callout(doc, "These figures must be yours",
            "Similarity, AI-detection, grammar and readability scores can only be produced by "
            "running the finished paper through the tools your institution uses (Turnitin, "
            "GPTZero, Grammarly or equivalent). They are left blank deliberately - a fabricated "
            "originality score would be the exact failure this case study is about. Run the "
            "scans on output/REVIEW_PAPER_IEEE.docx and record the results here.",
            fill="FDF3E7", border="D9A441")
    add_table(doc, pd.DataFrame([
        {"Parameter": "Similarity Index", "Required": "< 10%", "Result": "", "Tool used": "", "Date": ""},
        {"Parameter": "AI Detection Score", "Required": "< 10%", "Result": "", "Tool used": "", "Date": ""},
        {"Parameter": "Grammar Score", "Required": "> 90%", "Result": "", "Tool used": "", "Date": ""},
        {"Parameter": "Readability Score", "Required": "Good", "Result": "", "Tool used": "", "Date": ""},
    ]), widths=[1.6, 1.0, 1.2, 1.6, 1.0], font_pt=9)

    # ---- Phase 13
    doc.add_heading("Phase 13 - Final Review Paper Submission", level=2)
    add_prose(doc, "The complete review paper is submitted as a separate document, "
                   "**output/REVIEW_PAPER_IEEE.docx**, in two-column IEEE format with "
                   "numbered citations in order of first appearance.")
    add_table(doc, pd.DataFrame([{"#": i, "Section": s} for i, s in enumerate([
        "Title", "Abstract", "Keywords", "Introduction", "Research Methodology",
        "Literature Review", "Comparative Analysis", "Research Trends", "Research Gaps",
        "Future Research Directions", "Conclusion", "References"], start=1)]),
        widths=[0.5, 6], font_pt=9)

    doc.add_page_break()
    doc.add_heading("Appendix A - Faculty Evaluation Rubric", level=2)
    add_table(doc, pd.DataFrame([{"Criteria": c, "Marks": m} for c, m in [
        ("Topic Selection & Scope", 5), ("Paper Collection (30+)", 10),
        ("Literature Review Master Matrix", 15), ("Prompt Engineering Portfolio", 10),
        ("Research Evolution Analysis", 10), ("Comparative Study", 10),
        ("Limitation Mining", 10), ("Research Gap Identification", 10),
        ("Future Scope & Novel Ideas", 10), ("Humanization & Originality", 5),
        ("Final Review Paper Quality", 5), ("Total", 100)]]),
        widths=[4, 1], font_pt=9)

    doc.add_heading("Appendix B - Figures to Re-check Before Submission", level=2)
    cmp_check = load("comparison")
    flagged = cmp_check[cmp_check["figure_confidence"] == "check"]
    if flagged.empty:
        add_prose(doc, "No numeric figures are flagged for re-checking.")
    else:
        add_prose(doc, "Every numeric result quoted in this submission carries a confidence "
                       "flag. The rows below are figures recalled from reading rather than "
                       "re-read from the paper during compilation. Open each paper, confirm "
                       "the figure, and set figure_confidence to 'high' in comparison.csv.")
        add_table(doc, flagged[["id", "system", "headline"]].assign(Checked="[  ]")
                  .rename(columns={"id": "Paper", "system": "System",
                                   "headline": "Figure to confirm"}),
                  widths=[0.5, 1.5, 4.0, 0.7], font_pt=8)

    doc.add_heading("Appendix C - Reproducibility", level=2)
    add_table(doc, pd.DataFrame([
        {"File": "fetch_papers.py", "Purpose": "Resolves all 45 records against CrossRef and the arXiv API"},
        {"File": "seed_data.py", "Purpose": "Writes the Phase 3-11 analysis tables"},
        {"File": "content.py", "Purpose": "All prose and the full review paper text"},
        {"File": "build.py", "Purpose": "Renders the workbook and the IEEE paper; renumbers citations"},
        {"File": "data/papers.csv", "Purpose": "Verified bibliographic repository"},
        {"File": "data/matrix.csv", "Purpose": "Phase 3 master matrix, 35 papers x 8 fields"},
    ]), widths=[1.6, 5.4], font_pt=9)

    OUT.mkdir(exist_ok=True)
    path = OUT / "CASE_STUDY_2_WORKBOOK.docx"
    doc.save(str(path))
    return path


PUBLISHER_FAMILY = {
    "Association for Computational Linguistics": "ACL Anthology",
    "Institute of Electrical and Electronics Engineers (IEEE)": "IEEE",
    "MIT Press": "MIT Press (TACL)",
    "Springer Science and Business Media LLC": "Springer / Nature",
    "MDPI AG": "MDPI",
    "Elsevier BV": "Elsevier",
    "Wiley": "Wiley",
    "Association for Computing Machinery (ACM)": "ACM",
    "arXiv": "arXiv (preprint)",
}


def publisher_family(name: str) -> str:
    """Map raw CrossRef publisher strings onto the manual's source families."""
    if name in PUBLISHER_FAMILY:
        return PUBLISHER_FAMILY[name]
    for key, label in PUBLISHER_FAMILY.items():
        if key.split()[0].lower() in name.lower():
            return label
    return name


def strip_cites(text: str) -> str:
    """Workbook prose shows paper IDs directly rather than IEEE numerals."""
    return text


def check() -> int:
    papers, matrix = load("papers"), load("matrix")
    problems = []
    if len(matrix) < 30:
        problems.append(f"matrix has {len(matrix)} papers, minimum is 30")
    known = set(papers["id"])
    for name, col in [("matrix", "id"), ("limitations", "id"), ("comparison", "id")]:
        for pid in load(name)[col]:
            if pid not in known:
                problems.append(f"{name}.csv references unknown paper {pid}")
    for _, g in load("gaps").iterrows():
        for pid in re.findall(r"[PR]\d\d", g["support"]):
            if pid not in known:
                problems.append(f"gap {g['id']} cites unknown paper {pid}")
    dup = papers["doi"][papers["doi"] != ""].duplicated()
    if dup.any():
        problems.append("duplicate DOIs in papers.csv")

    if problems:
        print("CHECK FAILED")
        for p in problems:
            print("  -", p)
        return 1

    print("Check passed.")
    print(f"  papers        : {len(papers)} ({len(matrix)} in matrix, min 30)")
    print(f"  year range    : {papers['year'].min()}-{papers['year'].max()}")
    sections = {
        "abstract": content.ABSTRACT, "introduction": content.INTRODUCTION,
        "methodology": content.METHODOLOGY, "literature review": content.LITERATURE_REVIEW,
        "comparative": content.COMPARATIVE, "trends": content.TRENDS,
        "gaps intro": content.GAPS_INTRO, "conclusion": content.CONCLUSION,
    }
    total = sum(content.word_count(v) for v in sections.values())
    for k, v in sections.items():
        print(f"  {k:<18}: {content.word_count(v)} words")
    print(f"  narrative total   : {total} words (target 4000-6000 with tables/gaps expanded)")
    return 0


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--check", action="store_true")
    args = ap.parse_args()

    status = check()
    if status or args.check:
        return status

    print("\nBuilding review paper...")
    paper_docx, bib, words = build_paper()
    print(f"  {paper_docx.name}  ({words} words, {len(bib.order)} references cited)")
    uncited = bib.uncited()
    if uncited:
        print(f"  NOTE: {len(uncited)} collected papers not cited in the paper: {uncited}")

    print("\nBuilding workbook...")
    wb_docx = build_workbook(bib.order)
    print(f"  {wb_docx.name}")

    print("\nConverting to PDF...")
    for d in (paper_docx, wb_docx):
        pdf = to_pdf(d, d.with_suffix(".pdf"))
        print(f"  {pdf.name if pdf else d.stem + '.pdf SKIPPED (no converter)'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
